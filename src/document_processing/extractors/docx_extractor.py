# cockatoo_v1/src/document_processing/extractors/docx_extractor.py

"""
DOCX document extractor using python-docx.
"""

import os
import logging
from typing import Dict, Any, List, Union
from pathlib import Path
import zipfile

from .base_extractor import BaseExtractor

logger = logging.getLogger(__name__)

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. DOCX support will be limited.")


class DOCXExtractor(BaseExtractor):
    """
    Microsoft Word DOCX document extractor.
    """
    
    def __init__(self):
        """Initialize DOCX extractor."""
        super().__init__()
        if not HAS_DOCX:
            self.logger.warning(
                "python-docx is not installed. Install with: pip install python-docx"
            )
    
    def get_supported_formats(self) -> List[str]:
        """
        Return list of supported file formats.
        
        Returns:
            List of supported file extensions
        """
        return ['.docx', '.doc']
    
    def extract(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Dictionary containing:
                - text: Extracted text content
                - metadata: Document metadata
                - paragraphs: List of paragraphs with formatting info
                - tables: Extracted table data
                - images: Information about images in document
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        logger.info(f"Extracting DOCX: {file_path}")
        
        result = {
            "text": "",
            "metadata": {},
            "paragraphs": [],
            "tables": [],
            "images": [],
            "sections": [],
            "styles": {},
            "extraction_method": "python-docx" if HAS_DOCX else "zip_parsing",
        }
        
        # Add basic file metadata
        result["metadata"].update(self.get_basic_metadata(file_path))
        
        if HAS_DOCX:
            try:
                result = self._extract_with_docx(file_path, result)
            except Exception as e:
                logger.error(f"python-docx extraction failed: {e}")
                if file_path.suffix.lower() == '.docx':
                    logger.info("Falling back to zip parsing")
                    result = self._extract_with_zip(file_path, result)
        else:
            # Fallback to zip parsing
            if file_path.suffix.lower() == '.docx':
                result = self._extract_with_zip(file_path, result)
            else:
                raise ImportError(
                    "python-docx is required for .doc files. "
                    "Install with: pip install python-docx"
                )
        
        # Clean and post-process
        result["text"] = self.clean_text(result["text"])
        
        # Add language detection
        if result["text"]:
            result["metadata"]["language"] = self.detect_language(result["text"])
            result["metadata"]["summary"] = self.extract_summary(result["text"])
        
        return result
    
    def _extract_with_docx(self, file_path: Path, result: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract using python-docx library.
        
        Args:
            file_path: Path to DOCX file
            result: Result dictionary to update
            
        Returns:
            Updated result dictionary
        """
        doc = docx.Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraph_info = {
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "runs": [],
                }
                
                # Extract run information (for formatting)
                for run in para.runs:
                    if run.text.strip():
                        run_info = {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "font_size": run.font.size,
                            "font_name": run.font.name,
                        }
                        paragraph_info["runs"].append(run_info)
                
                paragraphs.append(paragraph_info)
                full_text.append(para.text)
        
        result["paragraphs"] = paragraphs
        result["text"] = "\n".join(full_text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                tables.append({
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                    "data": table_data,
                    "text": "\n".join([" | ".join(row) for row in table_data])
                })
        
        result["tables"] = tables
        
        # Extract metadata
        core_properties = doc.core_properties
        result["metadata"].update({
            "title": core_properties.title or "",
            "author": core_properties.author or "",
            "subject": core_properties.subject or "",
            "keywords": core_properties.keywords or "",
            "comments": core_properties.comments or "",
            "category": core_properties.category or "",
            "created": str(core_properties.created) if core_properties.created else "",
            "modified": str(core_properties.modified) if core_properties.modified else "",
            "last_modified_by": core_properties.last_modified_by or "",
            "revision": core_properties.revision or "",
            "version": core_properties.version or "",
        })
        
        # Extract sections information
        sections = []
        for section in doc.sections:
            section_info = {
                "start_type": str(section.start_type),
                "orientation": str(section.orientation),
                "page_width": section.page_width.inches if section.page_width else 0,
                "page_height": section.page_height.inches if section.page_height else 0,
            }
            sections.append(section_info)
        
        result["sections"] = sections
        
        # Count images (simplified)
        # Note: python-docx doesn't directly provide image content extraction
        rels = doc.part.rels
        image_count = 0
        for rel in rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        result["images_count"] = image_count
        
        return result
    
    def _extract_with_zip(self, file_path: Path, result: Dict[str, Any]) -> Dict[str, Any]:
        """
        Fallback extraction by parsing DOCX as ZIP archive.
        
        Args:
            file_path: Path to DOCX file
            result: Result dictionary to update
            
        Returns:
            Updated result dictionary
        """
        full_text = []
        
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Extract main document text
                if 'word/document.xml' in docx_zip.namelist():
                    with docx_zip.open('word/document.xml') as doc_file:
                        import re
                        content = doc_file.read().decode('utf-8', errors='ignore')
                        
                        # Simple XML parsing for text
                        text = re.sub(r'<[^>]+>', ' ', content)
                        text = re.sub(r'\s+', ' ', text).strip()
                        
                        full_text.append(text)
                
                # Try to extract metadata
                if 'docProps/core.xml' in docx_zip.namelist():
                    with docx_zip.open('docProps/core.xml') as core_file:
                        core_content = core_file.read().decode('utf-8', errors='ignore')
                        
                        # Extract metadata from core properties
                        metadata_patterns = {
                            'title': r'<dc:title[^>]*>(.*?)</dc:title>',
                            'creator': r'<dc:creator[^>]*>(.*?)</dc:creator>',
                            'subject': r'<dc:subject[^>]*>(.*?)</dc:subject>',
                            'description': r'<dc:description[^>]*>(.*?)</dc:description>',
                            'created': r'<dcterms:created[^>]*>(.*?)</dcterms:created>',
                            'modified': r'<dcterms:modified[^>]*>(.*?)</dcterms:modified>',
                        }
                        
                        for key, pattern in metadata_patterns.items():
                            match = re.search(pattern, core_content, re.DOTALL)
                            if match:
                                result["metadata"][key] = match.group(1).strip()
                
                result["text"] = " ".join(full_text)
                
        except zipfile.BadZipFile:
            logger.error(f"File is not a valid ZIP archive: {file_path}")
            result["text"] = ""
        except Exception as e:
            logger.error(f"Failed to parse DOCX as ZIP: {e}")
            result["text"] = ""
        
        return result
    
    def extract_structure(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extract document structure (headings, sections, etc.).
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Document structure information
        """
        if not HAS_DOCX:
            return {"error": "python-docx required for structure extraction"}
        
        file_path = Path(file_path)
        structure = {
            "headings": [],
            "sections": [],
            "lists": [],
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract headings
            for i, para in enumerate(doc.paragraphs):
                style_name = para.style.name.lower() if para.style else ""
                
                if 'heading' in style_name:
                    heading_level = 1
                    if 'heading 1' in style_name:
                        heading_level = 1
                    elif 'heading 2' in style_name:
                        heading_level = 2
                    elif 'heading 3' in style_name:
                        heading_level = 3
                    
                    structure["headings"].append({
                        "level": heading_level,
                        "text": para.text,
                        "paragraph_index": i,
                    })
            
            # Count lists
            list_count = 0
            for para in doc.paragraphs:
                if para._element.xpath('.//w:numPr'):
                    list_count += 1
            
            structure["lists"] = [{"count": list_count}]
            
        except Exception as e:
            logger.error(f"Failed to extract document structure: {e}")
        
        return structure